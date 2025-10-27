import streamlit as st
st.markdown(
    '''<style>
    body, .main, .block-container {
        background: #fff !important;
        color: #222 !important;
    }
    h1, h2, h3, h4, h5, h6 {
        color: #1976d2 !important;
        font-weight: bold !important;
        letter-spacing: 0.5px;
    }
    .stTextInput > div > input, .stTextArea textarea, .stSelectbox div[data-baseweb="select"] {
        background: #f7f7f7 !important;
        color: #222 !important;
        border: 2px solid #1976d2 !important;
        border-radius: 6px !important;
        font-size: 1.05em;
    }
    .stButton button {
        background: #1976d2 !important;
        color: #fff !important;
        border: 2px solid #ffd600 !important;
        border-radius: 6px !important;
        font-weight: bold;
        font-size: 1.08em;
        padding: 0.5em 1.2em;
        margin: 0.2em 0;
        transition: background 0.2s;
    }
    .stButton button:hover {
        background: #ffd600 !important;
        color: #1976d2 !important;
        border: 2px solid #1976d2 !important;
    }
    .stRadio div[role="radiogroup"] > div {
        background: #fff !important;
        color: #222 !important;
        border: 2px solid #1976d2 !important;
        border-radius: 6px !important;
        margin-bottom: 0.2em;
    }
    .stRadio label {
        color: #1976d2 !important;
        font-weight: bold;
    }
    /* Statusy WCAG */
    .stRadio div[role="radiogroup"] > div:nth-child(1) label { /* ✅ */
        background: #228B22 !important;
        color: #fff !important;
        border-radius: 4px;
        padding: 2px 10px;
    }
    .stRadio div[role="radiogroup"] > div:nth-child(2) label { /* ❌ */
        background: #C62828 !important;
        color: #fff !important;
        border-radius: 4px;
        padding: 2px 10px;
    }
    .stRadio div[role="radiogroup"] > div:nth-child(3) label { /* ⚠️ */
        background: #FFD600 !important;
        color: #222 !important;
        border-radius: 4px;
        padding: 2px 10px;
    }
    .stDownloadButton button {
        background: #ffd600 !important;
        color: #222 !important;
        border: 2px solid #1976d2 !important;
        border-radius: 6px !important;
        font-weight: bold;
        font-size: 1.08em;
        padding: 0.5em 1.2em;
        margin: 0.2em 0;
    }
    .stDownloadButton button:hover {
        background: #1976d2 !important;
        color: #fff !important;
        border: 2px solid #ffd600 !important;
    }
    .stTextArea textarea::placeholder, .stTextInput input::placeholder {
        color: #888 !important;
    }
    .stAlert, .stSuccess, .stError, .stWarning {
        border-radius: 6px !important;
        font-size: 1.05em;
    }
    </style>''', unsafe_allow_html=True)
from docx import Document
from datetime import datetime
import calendar
from openai_client import generate_recommendations
import os
import tempfile
import json
import shutil
import time
from docx.shared import Inches
from dotenv import load_dotenv

# Load .env if present
load_dotenv()

# Szablony kryteriów wg typu audytowanego pliku
wcag_templates = {
    "Dokument Word": [
        {
            "id": "1.1.1",
            "description": "Alternatywa w postaci tekstu dla treści nietekstowej",
            "long_description": "Każda treść nietekstowa (np. obraz, wykres, ikona) powinna mieć alternatywny opis tekstowy, który pozwala zrozumieć jej funkcję osobom korzystającym z czytników ekranu.",
            "example_test": "Sprawdź, czy wszystkie obrazy mają atrybut ALT opisujący ich funkcję. Otwórz dokument w czytniku ekranu i zweryfikuj, czy alternatywa jest czytelna i zrozumiała."
        },
        {
            "id": "1.2.1",
            "description": "Napisy dla mediów zmiennych w czasie (tylko audio)",
            "long_description": "Nagrania audio i wideo powinny mieć napisy lub transkrypcję, aby osoby niesłyszące mogły zrozumieć treść.",
            "example_test": "Odtwórz nagranie i sprawdź, czy dostępne są napisy lub transkrypcja."
        },
        {
            "id": "1.2.2",
            "description": "Napisy dla mediów zmiennych w czasie (wideo)",
            "long_description": "Nagrania wideo powinny mieć napisy rozszerzone, aby osoby niesłyszące mogły śledzić dialogi i dźwięki.",
            "example_test": "Odtwórz nagranie wideo i sprawdź, czy dostępne są napisy rozszerzone."
        },
        {
            "id": "1.2.4",
            "description": "Transkrypcja dla mediów zmiennych w czasie (wideo)",
            "long_description": "Treści wideo powinny mieć transkrypcję tekstową, aby osoby z niepełnosprawnościami mogły zapoznać się z treścią.",
            "example_test": "Sprawdź, czy do nagrania wideo dołączona jest transkrypcja tekstowa."
        },
        {
            "id": "1.3.1",
            "description": "Informacje i relacje",
            "long_description": "Struktura dokumentu powinna być logiczna, a relacje między elementami czytelne dla czytników ekranu.",
            "example_test": "Sprawdź nagłówki, listy, tabele – czy są poprawnie oznaczone i czytelne dla czytnika ekranu."
        },
        {
            "id": "1.3.2",
            "description": "Sekwencja sensowna",
            "long_description": "Kolejność prezentacji treści powinna być logiczna i zgodna z zamierzeniem autora.",
            "example_test": "Przejdź przez dokument tabulatorem i sprawdź, czy kolejność elementów jest sensowna."
        },
        {
            "id": "1.3.3",
            "description": "Charakterystyki sensoryczne",
            "long_description": "Instrukcje nie powinny opierać się wyłącznie na cechach zmysłowych (np. kolor, kształt).",
            "example_test": "Sprawdź, czy instrukcje nie odnoszą się tylko do koloru lub położenia, np. 'kliknij czerwony przycisk'."
        },
        {
            "id": "1.4.3",
            "description": "Kontrast (minimum)",
            "long_description": "Tekst i elementy graficzne powinny mieć odpowiedni kontrast względem tła, aby były czytelne dla osób słabowidzących.",
            "example_test": "Zmierz kontrast tekstu do tła (np. narzędziem online) – powinien wynosić co najmniej 4.5:1."
        },
        {
            "id": "1.4.4",
            "description": "Zmienność rozmiaru tekstu",
            "long_description": "Użytkownik powinien móc powiększyć tekst do 200% bez utraty czytelności lub funkcjonalności.",
            "example_test": "Powiększ tekst w przeglądarce/dokumencie i sprawdź, czy treść jest nadal czytelna."
        },
        {
            "id": "1.4.5",
            "description": "Obrazy tekstu",
            "long_description": "Tekst powinien być prezentowany jako tekst, nie jako obraz, chyba że jest to niezbędne (np. logo).",
            "example_test": "Sprawdź, czy teksty informacyjne nie są grafikami. Jeśli są, czy mają alternatywę tekstową."
        },
        {
            "id": "1.4.10",
            "description": "Przełączanie zawartości",
            "long_description": "Zawartość powinna być dostępna bez konieczności przewijania w poziomie na ekranach o szerokości 320px.",
            "example_test": "Otwórz dokument/stronę na wąskim ekranie i sprawdź, czy nie trzeba przewijać w poziomie."
        },
        {
            "id": "1.4.11",
            "description": "Kontrast elementów niestandardowych",
            "long_description": "Elementy interfejsu (np. przyciski, pola) powinny mieć kontrast co najmniej 3:1 względem tła.",
            "example_test": "Zmierz kontrast przycisków i pól formularza względem tła."
        },
        {
            "id": "1.4.12",
            "description": "Rozmiar tekstu w interfejsie",
            "long_description": "Odstępy między wierszami, akapitami, literami i wyrazami powinny być wystarczające dla czytelności.",
            "example_test": "Zmień odstępy w stylach CSS/dokumencie i sprawdź, czy treść jest czytelna."
        },
        {
            "id": "2.1.1",
            "description": "Obsługa klawiatury",
            "long_description": "Wszystkie funkcje powinny być dostępne z poziomu klawiatury, bez konieczności użycia myszy.",
            "example_test": "Przejdź przez dokument/stronę wyłącznie klawiaturą i sprawdź, czy można wykonać wszystkie operacje."
        },
        {
            "id": "2.1.2",
            "description": "Brak pułapek klawiaturowych",
            "long_description": "Fokus klawiatury nie powinien utknąć w żadnym elemencie – użytkownik musi móc go opuścić.",
            "example_test": "Przejdź tabulatorem przez wszystkie elementy i sprawdź, czy można opuścić każdy z nich."
        },
        {
            "id": "2.3.1",
            "description": "Trwałość błysków",
            "long_description": "Treści nie powinny zawierać elementów migających częściej niż 3 razy na sekundę, aby nie wywoływać ataków padaczki.",
            "example_test": "Sprawdź, czy na stronie/dokumencie nie ma animacji lub efektów migających powyżej 3 Hz."
        },
        {
            "id": "2.4.1",
            "description": "Możliwość pominięcia bloków",
            "long_description": "Użytkownik powinien mieć możliwość pominięcia powtarzających się bloków treści (np. menu).",
            "example_test": "Sprawdź, czy jest dostępny link lub mechanizm pomijania powtarzających się sekcji."
        },
        {
            "id": "2.4.3",
            "description": "Porządek fokusu",
            "long_description": "Kolejność przechodzenia fokusu powinna być logiczna i zgodna z zamierzeniem autora.",
            "example_test": "Przechodź tabulatorem przez elementy i sprawdź, czy kolejność jest sensowna."
        },
        {
            "id": "2.4.4",
            "description": "Cel linku (w kontekście)",
            "long_description": "Cel każdego linku powinien być jasny z samego tekstu lub kontekstu.",
            "example_test": "Sprawdź, czy linki mają zrozumiałe opisy, np. 'Więcej informacji' w kontekście."
        },
        {
            "id": "2.4.5",
            "description": "Wiele sposobów nawigacji",
            "long_description": "Powinno istnieć więcej niż jeden sposób na znalezienie strony lub sekcji (np. wyszukiwarka, mapa strony, menu).",
            "example_test": "Sprawdź, czy można dotrzeć do każdej sekcji na co najmniej dwa sposoby."
        },
        {
            "id": "2.4.6",
            "description": "Nagłówki i etykiety",
            "long_description": "Nagłówki i etykiety powinny jasno opisywać temat lub cel treści.",
            "example_test": "Sprawdź, czy nagłówki i etykiety są zrozumiałe i jednoznaczne."
        },
        {
            "id": "2.4.7",
            "description": "Widoczność fokusu",
            "long_description": "Fokus klawiatury powinien być zawsze widoczny podczas nawigacji.",
            "example_test": "Przechodź klawiaturą przez elementy i sprawdź, czy wskaźnik fokusu jest widoczny."
        },
        {
            "id": "3.1.1",
            "description": "Język strony",
            "long_description": "Język dokumentu/strony powinien być określony programowo, aby czytniki ekranu mogły go rozpoznać.",
            "example_test": "Sprawdź, czy w kodzie strony/dokumentu jest określony język (np. lang=\"pl\")."
        },
        {
            "id": "3.1.2",
            "description": "Język fragmentów",
            "long_description": "Język fragmentów tekstu różniących się od głównego powinien być oznaczony.",
            "example_test": "Sprawdź, czy cytaty, obce wyrażenia mają oznaczony język (np. lang=\"en\")."
        },
        {
            "id": "3.2.3",
            "description": "Spójność nawigacji",
            "long_description": "Powtarzające się mechanizmy nawigacji powinny być prezentowane w tej samej kolejności na wszystkich stronach.",
            "example_test": "Przejdź przez różne strony i sprawdź, czy menu, stopka są zawsze w tym samym miejscu."
        },
        {
            "id": "3.2.4",
            "description": "Spójność identyfikacji",
            "long_description": "Elementy o tej samej funkcji powinny być identyfikowane w ten sam sposób na wszystkich stronach.",
            "example_test": "Sprawdź, czy np. przycisk 'Wyślij' zawsze wygląda i działa tak samo."
        },
        {
            "id": "3.3.1",
            "description": "Identyfikacja błędów",
            "long_description": "Błędy wprowadzania danych powinny być jasno komunikowane użytkownikowi.",
            "example_test": "Wprowadź błędne dane w formularzu i sprawdź, czy pojawia się czytelny komunikat o błędzie."
        },
        {
            "id": "3.3.2",
            "description": "Sugestie dotyczące błędów",
            "long_description": "System powinien podpowiadać użytkownikowi, jak poprawić błędy w danych.",
            "example_test": "Wprowadź błędne dane i sprawdź, czy pojawia się sugestia poprawy."
        },
        {
            "id": "4.1.3",
            "description": "Status komunikatów",
            "long_description": "Komunikaty o stanie (np. sukces, błąd) powinny być dostępne dla technologii wspomagających.",
            "example_test": "Sprawdź, czy komunikaty o stanie są odczytywane przez czytnik ekranu."
        },
    ],
    "PDF": [
        {"id": "1.1.1", "description": "Alternatywa w postaci tekstu dla treści nietekstowej"},
        {"id": "1.2.1", "description": "Napisy dla mediów zmiennych w czasie (tylko audio)"},
        {"id": "1.2.2", "description": "Napisy dla mediów zmiennych w czasie (wideo)"},
        {"id": "1.2.4", "description": "Transkrypcja dla mediów zmiennych w czasie (wideo)"},
        {"id": "1.3.1", "description": "Informacje i relacje"},
        {"id": "1.3.2", "description": "Sekwencja sensowna"},
        {"id": "1.3.3", "description": "Charakterystyki sensoryczne"},
        {"id": "1.4.3", "description": "Kontrast (minimum)"},
        {"id": "1.4.4", "description": "Zmienność rozmiaru tekstu"},
        {"id": "1.4.5", "description": "Obrazy tekstu"},
        {"id": "1.4.10", "description": "Przełączanie zawartości"},
        {"id": "1.4.11", "description": "Kontrast elementów niestandardowych"},
        {"id": "1.4.12", "description": "Rozmiar tekstu w interfejsie"},
        {"id": "2.1.1", "description": "Obsługa klawiatury"},
        {"id": "2.1.2", "description": "Brak pułapek klawiaturowych"},
        {"id": "2.3.1", "description": "Trwałość błysków"},
        {"id": "2.4.1", "description": "Możliwość pominięcia bloków"},
        {"id": "2.4.3", "description": "Porządek fokusu"},
        {"id": "2.4.4", "description": "Cel linku (w kontekście)"},
        {"id": "2.4.5", "description": "Wiele sposobów nawigacji"},
        {"id": "2.4.6", "description": "Nagłówki i etykiety"},
        {"id": "2.4.7", "description": "Widoczność fokusu"},
        {"id": "3.1.1", "description": "Język strony"},
        {"id": "3.1.2", "description": "Język fragmentów"},
        {"id": "3.2.3", "description": "Spójność nawigacji"},
        {"id": "3.2.4", "description": "Spójność identyfikacji"},
        {"id": "3.3.1", "description": "Identyfikacja błędów"},
        {"id": "3.3.2", "description": "Sugestie dotyczące błędów"},
        {"id": "4.1.3", "description": "Status komunikatów"},
    ],
    "Strona WWW": [
        {"id": "1.1.1", "description": "Alternatywa w postaci tekstu dla treści nietekstowej"},
        {"id": "1.2.1", "description": "Napisy dla mediów zmiennych w czasie (tylko audio)"},
        {"id": "1.2.2", "description": "Napisy dla mediów zmiennych w czasie (wideo)"},
        {"id": "1.2.4", "description": "Transkrypcja dla mediów zmiennych w czasie (wideo)"},
        {"id": "1.3.1", "description": "Informacje i relacje"},
        {"id": "1.3.2", "description": "Sekwencja sensowna"},
        {"id": "1.3.3", "description": "Charakterystyki sensoryczne"},
        {"id": "1.4.3", "description": "Kontrast (minimum)"},
        {"id": "1.4.4", "description": "Zmienność rozmiaru tekstu"},
        {"id": "1.4.5", "description": "Obrazy tekstu"},
        {"id": "1.4.10", "description": "Przełączanie zawartości"},
        {"id": "1.4.11", "description": "Kontrast elementów niestandardowych"},
        {"id": "1.4.12", "description": "Rozmiar tekstu w interfejsie"},
        {"id": "2.1.1", "description": "Obsługa klawiatury"},
        {"id": "2.1.2", "description": "Brak pułapek klawiaturowych"},
        {"id": "2.3.1", "description": "Trwałość błysków"},
        {"id": "2.4.1", "description": "Możliwość pominięcia bloków"},
        {"id": "2.4.3", "description": "Porządek fokusu"},
        {"id": "2.4.4", "description": "Cel linku (w kontekście)"},
        {"id": "2.4.5", "description": "Wiele sposobów nawigacji"},
        {"id": "2.4.6", "description": "Nagłówki i etykiety"},
        {"id": "2.4.7", "description": "Widoczność fokusu"},
        {"id": "3.1.1", "description": "Język strony"},
        {"id": "3.1.2", "description": "Język fragmentów"},
        {"id": "3.2.3", "description": "Spójność nawigacji"},
        {"id": "3.2.4", "description": "Spójność identyfikacji"},
        {"id": "3.3.1", "description": "Identyfikacja błędów"},
        {"id": "3.3.2", "description": "Sugestie dotyczące błędów"},
        {"id": "4.1.3", "description": "Status komunikatów"},
    ],
}

st.title("AudytAI - Kreator raportów dostępności WCAG 2.1")


# Pola metadanych raportu
st.header("Dane audytu")
# Szkice: folder przechowujący robocze wersje raportów (docx + json)
draft_dir = os.path.join(os.getcwd(), "szkice")
os.makedirs(draft_dir, exist_ok=True)
draft_files = [f for f in os.listdir(draft_dir) if f.endswith('.json')]
selected_draft = st.selectbox("Wybierz szkic audytu", ["(brak)"] + draft_files, index=0, key="selected_draft")
if st.button("Wczytaj szkic") and selected_draft != "(brak)":
    try:
        with open(os.path.join(draft_dir, selected_draft), 'r', encoding='utf-8') as df:
            data = json.load(df)
        # Ustaw wartości w session_state przed renderowaniem widgetów
        # Ustaw wartości w session_state
        st.session_state['doc_type'] = data.get('doc_type', st.session_state.get('doc_type', 'Dokument Word'))
        st.session_state['doc_version'] = data.get('doc_version', '')
        st.session_state['audit_author'] = data.get('audit_author', '')
        st.session_state['browser_version'] = data.get('browser_version', '')
        st.session_state['app_name'] = data.get('app_name', '')
        st.session_state['tested_scope'] = data.get('tested_scope', '')
        # Load date into pickers
        try:
            d = datetime.fromisoformat(data.get('audit_date'))
            # Set pickers keys
            st.session_state['year_picker'] = d.year
            # month picker stores polish month name
            polish_months = [
                "styczeń", "luty", "marzec", "kwiecień", "maj", "czerwiec",
                "lipiec", "sierpień", "wrzesień", "październik", "listopad", "grudzień"
            ]
            st.session_state['month_picker'] = polish_months[d.month - 1]
            st.session_state['day_picker'] = d.day
        except Exception:
            pass
        # Responses and notes
        for k, v in data.get('responses', {}).items():
            st.session_state[k] = v
        for k, v in data.get('notes', {}).items():
            st.session_state[f"note_{k}"] = v
        st.rerun()
    except Exception as e:
        st.error(f"Błąd podczas wczytywania szkicu: {e}")

# Typ dokumentu — użyjemy key, żeby można było załadować szkic przez session_state
doc_type = st.selectbox("Wybierz typ pliku lub strony", ["Dokument Word", "PDF", "Strona WWW"], key='doc_type')
doc_version = st.text_input("Wersja dokumentu", value="", key='doc_version')
audit_author = st.text_input("Autor audytu", value="", key='audit_author')
browser_version = st.text_input("Wersja przeglądarki", value="", help="Np. Chrome 124.0.6367.61, Firefox 125.0.1", key='browser_version')
import locale
# Spróbuj kilku wariantów polskiej lokalizacji (Linux/macOS i Windows)
for loc in ("pl_PL.UTF-8", "pl_PL", "Polish_Poland.1250"):
    try:
        locale.setlocale(locale.LC_TIME, loc)
        break
    except Exception:
        continue
# Zamiast polegać na języku przeglądarki, wymuszamy picker z polskimi nazwami miesięcy
st.markdown("**Data audytu**")
today = datetime.today()
col_day, col_month, col_year = st.columns([1, 2, 1])
polish_months = [
    "styczeń", "luty", "marzec", "kwiecień", "maj", "czerwiec",
    "lipiec", "sierpień", "wrzesień", "październik", "listopad", "grudzień"
]
# Zakres lat: od -5 do +5 względem bieżącego roku



years = list(range(today.year - 5, today.year + 6))
# Domyślne wartości tylko raz na starcie
if 'year_picker' not in st.session_state:
    st.session_state['year_picker'] = today.year
if 'month_picker' not in st.session_state:
    st.session_state['month_picker'] = polish_months[today.month - 1]
if 'day_picker' not in st.session_state:
    st.session_state['day_picker'] = today.day
# Pobierz aktualne wartości z session_state
selected_year = st.session_state['year_picker']
selected_month_name = st.session_state['month_picker']
month_idx = polish_months.index(selected_month_name) + 1
last_day = calendar.monthrange(selected_year, month_idx)[1]
days = list(range(1, last_day + 1))
# Jeśli wybrany dzień jest poza zakresem, ustaw na ostatni dostępny dzień
if st.session_state['day_picker'] not in days:
    st.session_state['day_picker'] = last_day
selected_year_new = col_year.selectbox("Rok", years, index=years.index(selected_year))
selected_month_name_new = col_month.selectbox("Miesiąc", polish_months, index=polish_months.index(selected_month_name))
month_idx_new = polish_months.index(selected_month_name_new) + 1
last_day_new = calendar.monthrange(selected_year_new, month_idx_new)[1]
days_new = list(range(1, last_day_new + 1))
selected_day_new = col_day.selectbox("Dzień", days_new, index=min(days_new.index(st.session_state['day_picker']), len(days_new)-1))
# Aktualizuj session_state tylko jeśli użytkownik zmienił wybór
if selected_year_new != st.session_state['year_picker']:
    st.session_state['year_picker'] = selected_year_new
if selected_month_name_new != st.session_state['month_picker']:
    st.session_state['month_picker'] = selected_month_name_new
if selected_day_new != st.session_state['day_picker']:
    st.session_state['day_picker'] = selected_day_new
# Zbuduj obiekt datetime na podstawie wyboru
try:
    audit_date = datetime(selected_year, month_idx, selected_day)
except Exception:
    # Bezpieczny fallback na dziś
    audit_date = datetime.today()


# Ręczne formatowanie polskiej daty (odmienione miesiące)
polish_months_gen = [
    "stycznia", "lutego", "marca", "kwietnia", "maja", "czerwca",
    "lipca", "sierpnia", "września", "października", "listopada", "grudnia"
]
def format_polish_date(dt):
    return f"{dt.day} {polish_months_gen[dt.month-1]} {dt.year}"

formatted = format_polish_date(audit_date)
st.markdown(f"**Wybrana data audytu:** {formatted}")

# ustaw kryteria na podstawie wybranego typu dokumentu
wcag_criteria = wcag_templates.get(doc_type, wcag_templates["Dokument Word"]) 

st.header("Lista wytycznych WCAG")
responses = {}

# Metadane raportu: nazwa aplikacji/dokumentu i zakres stron/testów
app_name = st.text_input("Nazwa aplikacji / dokumentu", value="", help="Np. 'Strona główna MySite' lub 'Dokument X'", key='app_name')
tested_scope = st.text_area("Zakres testu (np. lista stron, sekcji)", value="", help="Wypisz adresy URL lub numery stron, które były testowane, oddzielone nową linią", key='tested_scope')

# Dodaj przycisk zapisu do szkicu

# Wybór modelu (domyślnie najtańsza opcja)
model_choice = st.selectbox(
    "Wybierz model (domyślnie MOCK)",
    ("gpt-3.5-turbo", "gpt-4", "MOCK"),
    index=2,
    help="gpt-3.5-turbo — najtańsza opcja; MOCK — tryb developerski bez wywołań OpenAI",
    key='model_choice'
)

# Debug: pokaż wygenerowane rekomendacje w UI (przed zapisem do pliku)
debug_show_recs = st.checkbox(
    "Pokaż rekomendacje (AI) w aplikacji po generowaniu",
    value=True,
    help="Zaznacz, aby zobaczyć tekst rekomendacji w UI (przydatne do debugowania)")

notes = {}
uploads = {}
for crit in wcag_criteria:
    c1, c2 = st.columns([2, 5])
    # Wytyczna z większą czcionką
    c1.markdown(f"<span style='font-size:1.2em;font-weight:bold'>{crit['id']} - {crit['description']}</span>", unsafe_allow_html=True)
    responses[crit["id"]] = c1.radio(
        "Status wytycznej:",
        ("✅ Spełnione", "❌ Niespełnione", "⚠️ Nie dotyczy"),
        key=crit["id"]
    )
    # Modal info button
    show_info = c1.button(f"Co to znaczy?", key=f"info_{crit['id']}")
    if show_info:
        st.session_state[f"modal_{crit['id']}"] = True
    # Modal close logic: handle before rendering modal content
    modal_key = f"modal_{crit['id']}"
    close_key = f"close_{crit['id']}"
    # Obsługa zamknięcia modala na początku pętli
    if st.session_state.get(modal_key, False) and st.session_state.get(close_key, False):
        st.session_state[modal_key] = False
        st.session_state[close_key] = False

    # Renderowanie modala
    if st.session_state.get(modal_key, False):
        st.markdown(f"### {crit['id']} - {crit['description']}")
        st.info(f"Opis: {crit.get('long_description', crit.get('description', ''))}")
        example = crit.get('example_test', None)
        if example:
            st.markdown(f"**Jak testować:** {example}")
        st.markdown("<br>", unsafe_allow_html=True)
        st.button("Zamknij", key=close_key)
    # Notatka i upload w jednej kolumnie, labelki powiększone (zawsze widoczne)
    required = responses.get(crit["id"]) == "❌ Niespełnione"
    # Use full phrase "Notatka do wytycznej <id>" for clarity
    label = f"<span style='font-size:1.05em;font-weight:bold'>Notatka do wytycznej {crit['id']}{' <span style=\'color:red\'>*</span>' if required else ''}</span>"
    c2.markdown(label, unsafe_allow_html=True)
    notes[crit["id"]] = c2.text_area(
        f"Notatka do wytycznej {crit['id']}", value="", key=f"note_{crit['id']}", height=180,
        help="Wklej opis, kontekst, kroki reprodukcji. Możesz też wkleić linki lub krótki kod.",
        placeholder=f"Notatka do {crit['id']}: opisz problem, kroki reprodukcji, kontekst (wymagane jeśli kryterium nie jest niespełnione)."
    )
    # Try to use the newer 'button_label' arg when available; fall back if Streamlit version is older
    try:
        uploads[crit["id"]] = c2.file_uploader(
            "Wybierz pliki (zrzuty ekranu, zdjęcia)", type=["png", "jpg", "jpeg"], key=f"img_{crit['id']}",
            label_visibility="visible",
            help="Dodaj pliki, które pokazują problem lub potwierdzają spełnienie wytycznej.",
            accept_multiple_files=True,
            button_label="Dodaj plik"
        )
    except TypeError:
        # Older Streamlit versions don't support button_label
        uploads[crit["id"]] = c2.file_uploader(
            "Wybierz pliki (zrzuty ekranu, zdjęcia)", type=["png", "jpg", "jpeg"], key=f"img_{crit['id']}",
            label_visibility="visible",
            help="Dodaj pliki, które pokazują problem lub potwierdzają spełnienie wytycznej.",
            accept_multiple_files=True
        )
import pandas as pd

if 'report_ready' not in st.session_state:
    st.session_state['report_ready'] = False
if 'report_filename' not in st.session_state:
    st.session_state['report_filename'] = None

if st.button("Generuj raport Word"):
    # Wygeneruj rekomendacje AI przed budową raportu Word
    recs = generate_recommendations(responses, notes, model_choice)
    # Walidacja: wymagalność notatki dla niespełnionych kryteriów
    missing_notes = []
    for crit in wcag_criteria:
        if responses.get(crit["id"]) == "❌ Niespełnione" and not notes.get(crit["id"]):
            missing_notes.append(crit["id"])
    if missing_notes:
        st.warning(f"Dla wytycznych oznaczonych jako niespełnione wymagane jest uzupełnienie notatki: {', '.join(missing_notes)}")
    else:
        # Funkcja czyszcząca rekomendacje AI
        import re
        def format_recommendation(rec):
            # rec może być dict lub string
            if not rec:
                return ""
            import re
            # Funkcja czyszcząca tekst z multimediów i niedozwolonych linków
            def strip_multimedia(text):
                # Usuń tagi multimedialne
                text = re.sub(r'<\s*(video|source|img|iframe|audio|embed)[^>]*>', '', text, flags=re.IGNORECASE)
                # Usuń linki do wideo
                text = re.sub(r'https?://[^\s]+(youtube|vimeo|\.mp4|\.webm|\.mov|\.avi|\.wmv|\.mkv)[^\s]*', '', text, flags=re.IGNORECASE)
                # Usuń komunikaty o nieobsługiwanym wideo
                text = re.sub(r'Nieobsługiwany format wideo|Nie można odtworzyć wideo|Brak podglądu', '', text, flags=re.IGNORECASE)
                return text.strip()

            if isinstance(rec, dict):
                parts = []
                # Numer kryterium lub nazwa (jeśli dostępna)
                wcag_id = rec.get('id') or rec.get('nazwa')
                if wcag_id:
                    wcag_id = strip_multimedia(str(wcag_id))
                    if wcag_id:
                        parts.append(f"{wcag_id}")
                # Rekomendacja
                naprawa = rec.get('naprawa') or rec.get('rekomendacja')
                if naprawa:
                    naprawa = strip_multimedia(str(naprawa).strip().strip(':').strip('"'))
                    if naprawa:
                        parts.append(f"Rekomendacja: {naprawa}")
                # Przykład poprawy
                kod = rec.get('kod') or rec.get('przyklad')
                if kod:
                    kod = strip_multimedia(str(kod).strip().strip(':').strip('"'))
                    if kod:
                        parts.append(f"Przykład poprawy: {kod}")
                # Źródło/W3C link
                w3c_link = rec.get('w3c_link') or rec.get('źródło') or rec.get('zrodlo')
                if w3c_link:
                    w3c_link = strip_multimedia(str(w3c_link).strip().strip(':').strip('"'))
                    if w3c_link:
                        parts.append(f"Źródło: {w3c_link}")
                # Usuń puste linie
                return "\n".join([p for p in parts if p])
            # Jeśli rec to string, wyczyść i wypisz jako rekomendację
            clean = strip_multimedia(str(rec).strip().strip(':').strip('"'))
            if clean:
                return f"Rekomendacja: {clean}"
            return ""

        # Debug: pokaż rekomendacje w UI jeśli checkbox zaznaczony
        if debug_show_recs and recs:
            st.subheader("Rekomendacje AI (podgląd)")
            def render_bold_recommendation(formatted, cid=None):
                lines = formatted.split("\n")
                html = ""
                if cid:
                    html += f"<b>{cid}</b><br>"
                for line in lines:
                    if line.startswith("Rekomendacja:"):
                        html += f"<b>Rekomendacja:</b> {line[len('Rekomendacja:'):].lstrip()}<br>"
                    elif line.startswith("Źródło:"):
                        html += f"<br><b>Źródło:</b> {line[len('Źródło:'):].lstrip()}<br>"
                    else:
                        html += line + "<br>"
                st.markdown(html, unsafe_allow_html=True)

            if isinstance(recs, dict):
                for cid, rec in recs.items():
                    formatted = format_recommendation(rec)
                    if formatted:
                        render_bold_recommendation(formatted, cid)
            else:
                formatted = format_recommendation(recs)
                if formatted:
                    render_bold_recommendation(formatted)
        from docx.shared import Pt, RGBColor, Cm
        doc = Document()
        # Ustaw czcionkę globalnie na Calibri 11pt
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Tytuł główny
        if app_name:
            title = f"Raport z audytu dostępności WCAG 2.1 - {app_name}"
        else:
            title = "Raport z audytu dostępności WCAG 2.1"
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(16)
        p.alignment = 1  # center

        # Podtytuł
        if doc_type:
            p2 = doc.add_paragraph()
            run2 = p2.add_run(f"Typ audytowanego dokumentu: {doc_type}")
            run2.italic = True
            p2.alignment = 1

        # Wprowadzenie
        intro = doc.add_paragraph()
        intro.add_run("Niniejszy raport przedstawia wyniki audytu dostępności cyfrowej zgodnie z wytycznymi WCAG 2.1. Celem audytu jest wskazanie mocnych stron oraz obszarów wymagających poprawy, aby zapewnić pełną dostępność dla wszystkich użytkowników. Poniżej przedstawiono szczegółowe rekomendacje oraz status spełnienia poszczególnych kryteriów.")
        doc.add_paragraph("")

        # Metadane
        meta_items = [
            ("Data audytu", format_polish_date(audit_date)),
            ("Autor audytu", audit_author),
            ("Wersja dokumentu", doc_version),
            ("Nazwa aplikacji / dokumentu", app_name),
            ("Wersja przeglądarki", browser_version),
            ("Zakres testu", tested_scope.replace('\n', ', '))
        ]
        for label, value in meta_items:
            if value:
                para = doc.add_paragraph()
                run_label = para.add_run(f"{label}: ")
                run_label.bold = True
                para.add_run(str(value))
        doc.add_paragraph("")

        # Tabela rekomendacji
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        headers = ['ID / Nazwa', 'Opis', 'Status', 'Uwagi']
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        # Ustaw szerokości kolumn
        try:
            hdr_cells[0].width = Cm(2)
            hdr_cells[1].width = Cm(7)
            hdr_cells[2].width = Cm(3)
            hdr_cells[3].width = Cm(5)
        except Exception:
            pass

        # Kolory statusów
        status_colors = {
            "Spełnione": RGBColor(198, 239, 206),   # jasny zielony
            "Niespełnione": RGBColor(255, 199, 206), # jasny czerwony
            "Nie dotyczy": RGBColor(255, 235, 156)   # pastelowy żółty
        }

        for crit in wcag_criteria:
            row_cells = table.add_row().cells
            # ID / Nazwa
            row_cells[0].text = f"{crit['id']} / {crit['description']}"
            # Opis
            row_cells[1].text = crit.get('long_description', crit['description'])
            # Status
            status = responses[crit['id']]
            status_txt = status.replace('✅ ', '').replace('❌ ', '').replace('⚠️ ', '')
            row_cells[2].text = status_txt
            # Kolor tła komórki statusu (poprawnie przez XML)
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls
            color_map = {
                "Spełnione": "C6EFCE",   # jasny zielony
                "Niespełnione": "FFC7CE", # jasny czerwony
                "Nie dotyczy": "FFEB9C"   # pastelowy żółty
            }
            fillcolor = color_map.get(status_txt, "FFFFFF")
            tc = row_cells[2]._tc
            tcPr = tc.get_or_add_tcPr()
            tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fillcolor}"/>'))
            # Uwagi
            note = notes.get(crit['id'], "")
            row_cells[3].text = note
            # Czcionka w całym wierszu
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(11)
            # Wyrównanie tekstu do lewej
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = 0

        # Odstęp po tabeli
        doc.add_paragraph("")
        # Dodaj sekcję rekomendacji AI na końcu
        doc.add_page_break()
        # Dodaj sekcję rekomendacji AI tylko jeśli recs istnieje i nie jest None
        if 'recs' in locals() and recs is not None:
            doc.add_heading("Rekomendacje AI dla niespełnionych kryteriów", level=2)
            for crit in wcag_criteria:
                cid = crit["id"]
                rec = None
                if isinstance(recs, dict):
                    rec = recs.get(cid) or recs.get("_combined")
                else:
                    rec = None
                formatted = format_recommendation(rec)
                if formatted:
                    doc.add_heading(f"{cid} - {crit['description']}", level=3)
                    # Lepsze formatowanie: osobne akapity dla etykiet i treści
                    lines = formatted.split("\n")
                    for line in lines:
                        if line.startswith("Rekomendacja:"):
                            para_label = doc.add_paragraph()
                            run_label = para_label.add_run("Rekomendacja:")
                            run_label.bold = True
                            para_text = doc.add_paragraph(line[len("Rekomendacja:"):].lstrip())
                        elif line.startswith("Źródło:"):
                            para_label = doc.add_paragraph()
                            run_label = para_label.add_run("Źródło:")
                            run_label.bold = True
                            para_text = doc.add_paragraph(line[len("Źródło:"):].lstrip())
                        else:
                            doc.add_paragraph(line)
        # ... rekomendacje AI i dalsza logika ...
        # Zapis dokumentu
        def _slugify(name: str) -> str:
            import re
            s = name.strip().lower()
            s = re.sub(r"[^a-z0-9]+", "_", s)
            s = re.sub(r"_+", "_", s).strip("_")
            return s or "report"
        date_str = datetime.today().strftime('%Y-%m-%d')
        if app_name:
            safe = _slugify(app_name)
            filename = f"Raport_WCAG_{safe}_{date_str}.docx"
        else:
            filename = f"Raport_WCAG_{date_str}.docx"
        doc.save(filename)
        st.session_state['report_ready'] = True
        st.session_state['report_filename'] = filename
        # Automatyczny zapis szkicu po wygenerowaniu raportu Word
        date_str = datetime.today().strftime('%Y-%m-%d')
        draft_dir = os.path.join(os.getcwd(), "szkice")
        os.makedirs(draft_dir, exist_ok=True)
        draft_data = {
            "doc_type": doc_type,
            "doc_version": doc_version,
            "audit_author": audit_author,
            "browser_version": browser_version,
            "app_name": app_name,
            "tested_scope": tested_scope,
            "audit_date": audit_date.isoformat(),
            "responses": responses,
            "notes": notes
        }
        safe = app_name.strip().lower().replace(" ", "_") or "szkic"
        draft_name = f"{safe}_{date_str}.json"
        draft_path = os.path.join(draft_dir, draft_name)
        try:
            with open(draft_path, "w", encoding="utf-8") as df:
                json.dump(draft_data, df, ensure_ascii=False, indent=2)
            st.success(f"Szkic zapisany: {draft_name}")
        except Exception as e:
            st.error(f"Błąd zapisu szkicu: {e}")

# Przycisk pobierania raportu zawsze widoczny, jeśli raport został wygenerowany
if st.session_state.get('report_ready') and st.session_state.get('report_filename'):
    try:
        filesize = None
        with open(st.session_state['report_filename'], "rb") as f:
            data = f.read()
            filesize = len(data)
            st.download_button("Pobierz raport (Word)", data=data, file_name=st.session_state['report_filename'], mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="download_word")
        st.success(f"Raport wygenerowany: {st.session_state['report_filename']} ({filesize} bajtów)")
        # Automatyczna konwersja do PDF
        try:
            from docx2pdf import convert
            pdf_name = st.session_state['report_filename'].replace('.docx', '.pdf')
            convert(st.session_state['report_filename'], pdf_name)
            import os
            if os.path.exists(pdf_name) and os.path.getsize(pdf_name) > 0:
                with open(pdf_name, "rb") as fpdf:
                    st.download_button("Pobierz raport (PDF)", data=fpdf.read(), file_name=pdf_name, mime="application/pdf", key="download_pdf")
                st.success(f"Raport PDF zapisany: {pdf_name}")
            else:
                st.error(f"Nie udało się wygenerować PDF: plik nie powstał lub jest pusty.")
        except ImportError:
            st.warning("Aby eksportować do PDF, zainstaluj pakiet docx2pdf: pip install docx2pdf")
        except Exception as e:
            import os
            pdf_name = st.session_state['report_filename'].replace('.docx', '.pdf')
            if os.path.exists(pdf_name) and os.path.getsize(pdf_name) > 0:
                with open(pdf_name, "rb") as fpdf:
                    st.download_button("Pobierz raport (PDF)", data=fpdf.read(), file_name=pdf_name, mime="application/pdf", key="download_pdf")
                st.success(f"Raport PDF zapisany: {pdf_name}")
                # Nie pokazuj ostrzeżenia jeśli PDF jest poprawny
            else:
                st.error(f"Nie udało się wygenerować PDF: {e}")
    except Exception as e:
        st.error(f"Utworzono plik, ale nie udało się przygotować przycisku pobierania: {e}")

    # (zapis pliku Word odbywa się wyłącznie podczas generowania raportu, nie powielaj tutaj)

    # cleanup temp image files
    try:
        for tf in temp_files:
            if os.path.exists(tf):
                os.unlink(tf)
    except Exception:
        pass

    # Zapisz metadane do Excela
    meta = {
        "Data audytu": audit_date.strftime('%Y-%m-%d'),
        "Autor audytu": audit_author,
        "Wersja dokumentu": doc_version,
        "Nazwa aplikacji / dokumentu": app_name,
        "Typ audytowanego dokumentu": doc_type,
        "Wersja przeglądarki": browser_version,
        "Zakres testu": tested_scope.replace('\n', ', ')
    }

    import pandas as pd
    meta_df = pd.DataFrame([meta])
    xlsx_name = st.session_state['report_filename'].replace('.docx', '.xlsx')
    meta_df.to_excel(xlsx_name, index=False)
    try:
        with open(xlsx_name, "rb") as f:
            st.download_button("Pobierz metadane (Excel)", data=f.read(), file_name=xlsx_name, mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        st.success(f"Metadane zapisane: {xlsx_name}")
        # Zapisz szkic automatycznie po pobraniu metadanych
        def save_draft():
            draft_dir = os.path.join(os.getcwd(), "szkice")
            os.makedirs(draft_dir, exist_ok=True)
            draft_data = {
                "doc_type": doc_type,
                "doc_version": doc_version,
                "audit_author": audit_author,
                "browser_version": browser_version,
                "app_name": app_name,
                "tested_scope": tested_scope,
                "audit_date": audit_date.isoformat(),
                "responses": responses,
                "notes": notes
            }
            date_str = datetime.today().strftime('%Y-%m-%d')
            safe = app_name.strip().lower().replace(" ", "_") or "szkic"
            draft_name = f"{safe}_{date_str}.json"
            draft_path = os.path.join(draft_dir, draft_name)
            with open(draft_path, "w", encoding="utf-8") as df:
                json.dump(draft_data, df, ensure_ascii=False, indent=2)
        save_draft()
    except Exception as e:
        st.error(f"Utworzono plik Excel, ale nie udało się przygotować przycisku pobierania: {e}")
